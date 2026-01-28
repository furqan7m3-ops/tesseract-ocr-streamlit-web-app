"""Word document generation module."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


class WordDocumentGenerator:
    """Generate Word documents from OCR results with formatting."""
    
    def __init__(self, output_path="output.docx"):
        self.output_path = output_path
        self.document = Document()
    
    def add_title(self, title, font_size=14, is_bold=True):
        """Add title to document."""
        p = self.document.add_paragraph(title)
        run = p.runs[0]
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return p
    
    def add_extracted_text(self, text, alignment='left', formatting_info=None):
        """
        Add extracted text with formatting.
        
        Args:
            text: Text to add
            alignment: 'left', 'center', or 'right'
            formatting_info: Dict with bold, italic, font_size info
        """
        p = self.document.add_paragraph(text)
        
        # Set alignment
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        p.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)
        
        # Apply formatting if provided
        if formatting_info and p.runs:
            run = p.runs[0]
            
            if formatting_info.get('is_bold'):
                run.font.bold = True
            
            if formatting_info.get('is_italic'):
                run.font.italic = True
            
            if 'font_size' in formatting_info:
                run.font.size = Pt(formatting_info['font_size'])
        
        return p
    
    def add_text_blocks(self, text_blocks, formatting_info=None):
        """Add text blocks organized into paragraphs."""
        if not text_blocks:
            return
        
        # If text_blocks is a list of dictionaries (paragraph blocks)
        for block_group in text_blocks:
            paragraph_text = []
            
            if isinstance(block_group, list):
                # Group of text blocks forming a paragraph
                for block in block_group:
                    if isinstance(block, dict) and 'text' in block:
                        paragraph_text.append(block['text'])
            else:
                # Single block
                paragraph_text.append(str(block_group))
            
            if paragraph_text:
                full_text = ' '.join(paragraph_text)
                alignment = formatting_info.get('alignment', 'left') if formatting_info else 'left'
                self.add_extracted_text(full_text, alignment=alignment, 
                                       formatting_info=formatting_info)
    
    def add_raw_text(self, text):
        """Add raw text without special formatting."""
        self.document.add_paragraph(text)
    
    def add_image(self, image_path, width=Inches(5)):
        """Add image to document."""
        if os.path.exists(image_path):
            self.document.add_picture(image_path, width=width)
    
    def add_heading(self, text, level=1):
        """Add heading to document."""
        self.document.add_heading(text, level=level)
    
    def add_page_break(self):
        """Add page break."""
        self.document.add_page_break()
    
    def save(self, output_path=None):
        """Save document to file."""
        if output_path:
            self.output_path = output_path
        
        self.document.save(self.output_path)
        return self.output_path
    
    def set_metadata(self, title="", author="", subject=""):
        """Set document metadata."""
        core_props = self.document.core_properties
        if title:
            core_props.title = title
        if author:
            core_props.author = author
        if subject:
            core_props.subject = subject


def create_ocr_document(text, image_path=None, formatting_info=None, output_path="output.docx"):
    """
    Convenience function to create OCR document with extracted text.
    
    Args:
        text: Extracted text
        image_path: Path to original image (optional)
        formatting_info: Dictionary with formatting details
        output_path: Output Word document path
    
    Returns:
        str: Path to created document
    """
    doc_gen = WordDocumentGenerator(output_path)
    
    # Add title
    doc_gen.add_title("OCR Extracted Document", font_size=16)
    
    # Add original image if provided
    if image_path and os.path.exists(image_path):
        doc_gen.add_heading("Source Image", level=2)
        doc_gen.add_image(image_path)
        doc_gen.add_page_break()
    
    # Add extracted text
    doc_gen.add_heading("Extracted Text", level=2)
    
    if isinstance(text, dict) and 'text_blocks' in text:
        # Text with formatting information
        doc_gen.add_text_blocks(text['text_blocks'], formatting_info=formatting_info or text.get('formatting'))
    else:
        # Plain text
        doc_gen.add_raw_text(str(text))
    
    # Save and return path
    return doc_gen.save()
